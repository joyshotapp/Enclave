"""
Phase 11 — 內容生成引擎 (Content Generator)

基於知識庫文件，生成各種結構化內容草稿：
  - 單案件文件草稿（答辯狀、函件、摘要）
  - 跨案件分析報告（案件類型分析、趨勢整理）
  - 會議記錄整理（上傳錄音/筆記 → 生成結構化摘要）
  - FAQ 草稿（從高頻問答自動生成）

設計原則：
  - 每段生成內容必須標注來源文件（引用面板）
  - 生成結果可匯出 Word（python-docx）或 PDF（reportlab）
  - 生成過程使用 SSE 串流，讓使用者看到即時進度
  - LLM Provider 抽象層（OpenAI / Ollama 可切換）

生成模板（可擴充）：
  DRAFT_RESPONSE     — 答覆/函件草稿
  CASE_SUMMARY       — 案件摘要
  MEETING_MINUTES    — 會議記錄
  ANALYSIS_REPORT    — 分析報告
  FAQ_DRAFT          — FAQ 草稿

TODO (Phase 11 實作):
  - 設計各模板的 system prompt
  - 實作 RAG 增強生成（先檢索相關文件，再生成）
  - SSE 串流輸出
  - Word/PDF 匯出功能
  - 來源標注（citation tracking）
"""

import logging
from enum import Enum
from typing import AsyncIterator, List, Optional

logger = logging.getLogger(__name__)


class GenerationTemplate(str, Enum):
    DRAFT_RESPONSE = "draft_response"       # 函件/答覆草稿
    CASE_SUMMARY = "case_summary"           # 案件摘要
    MEETING_MINUTES = "meeting_minutes"     # 會議記錄整理
    ANALYSIS_REPORT = "analysis_report"     # 跨案件分析報告
    FAQ_DRAFT = "faq_draft"                 # FAQ 草稿


class ContentGenerator:
    """RAG 增強的內容生成引擎。"""

    def __init__(self, llm_client=None, retriever=None):
        self.llm_client = llm_client
        self.retriever = retriever

    async def generate_stream(
        self,
        template: GenerationTemplate,
        user_prompt: str,
        context_query: str,
        tenant_id: str,
        max_tokens: int = 3000,
        extra_context: str = "",  # P11-3: pre-fetched document content injected by caller
    ) -> AsyncIterator[str]:
        """
        RAG 增強生成，以 SSE 串流回傳。

        流程：
          1. 用 context_query 從知識庫檢索相關文件
          2. 組合 system prompt（依 template）+ 相關文件 + user_prompt
          3. 呼叫 LLM，串流回傳生成內容
          4. 在最後附上引用來源清單
        """
        import asyncio
        from uuid import UUID

        # ── 1. 知識庫檢索 ───────────────────────────────────────────
        sources: List[dict] = []
        context_text = ""

        if context_query and self.retriever:
            try:
                loop = asyncio.get_event_loop()
                results = await loop.run_in_executor(
                    None,
                    lambda: self.retriever.search(
                        tenant_id=UUID(tenant_id),
                        query=context_query,
                        top_k=5,
                    ),
                )
                sources = results or []
                if sources:
                    parts: List[str] = []
                    for i, r in enumerate(sources, 1):
                        doc_name = (
                            r.get("document_name")
                            or r.get("filename")
                            or f"文件{i}"
                        )
                        text = r.get("content") or r.get("text") or ""
                        parts.append(f"【文件 {i}：{doc_name}】\n{text[:1500]}")
                    context_text = "\n\n".join(parts)
            except Exception as exc:
                logger.warning("KB retrieval failed: %s", exc)

        # ── 2. 組合 system prompt ────────────────────────────────────
        system_prompt = self._get_system_prompt(template)

        # P11-3: Merge extra_context (direct documents) + KB retrieval context
        full_context = ""
        if extra_context:
            full_context += f"【直接引用文件】\n{extra_context}\n\n"
        if context_text:
            full_context += f"【知識庫檢索結果】\n{context_text}"

        if full_context:
            system_prompt += (
                "\n\n以下是可供參考的文件資料，請根據這些文件生成內容。"
                "\n**重要：在行文中，當你引用了某份文件的資訊時，請在句末加上 [來源 N] 標記"
                "（N 為下方來源編號），以便讀者追溯原始資料。**"
                f"\n\n{full_context}"
            )

        # ── 3. LLM 串流生成 ──────────────────────────────────────────
        if self.llm_client:
            async for chunk in self.llm_client.stream(
                system_prompt, user_prompt, max_tokens=max_tokens
            ):
                yield chunk
        else:
            yield "⚠️ LLM 服務目前不可用，請檢查環境設定中的 OPENAI_API_KEY 或 OLLAMA_BASE_URL。"
            return

        # ── 4. 附上引用來源 ─────────────────────────────────────────
        if sources:
            yield "\n\n---\n**引用來源**\n"
            for i, r in enumerate(sources, 1):
                doc_name = (
                    r.get("document_name")
                    or r.get("filename")
                    or f"文件{i}"
                )
                score = r.get("score") or r.get("similarity")
                score_str = f"（相似度 {score:.2f}）" if isinstance(score, float) else ""
                yield f"{i}. {doc_name}{score_str}\n"

    async def export_to_docx(self, content: str, title: str, sources: List[dict]) -> bytes:
        """將生成內容匯出為 Word 檔案，回傳 bytes。"""
        from docx import Document as DocxDocument  # type: ignore
        from docx.shared import Pt  # type: ignore
        from io import BytesIO

        doc = DocxDocument()

        # 標題
        doc.add_heading(title, level=1)

        # 內容（依 Markdown 標題層級拆段）
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("---"):
                doc.add_paragraph("─" * 50)
            else:
                doc.add_paragraph(stripped)

        # 引用來源
        if sources:
            doc.add_heading("引用來源", level=2)
            for i, s in enumerate(sources, 1):
                doc_name = (
                    s.get("document_name")
                    or s.get("filename")
                    or f"文件{i}"
                )
                doc.add_paragraph(f"{i}. {doc_name}", style="List Number")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def export_to_pdf(self, content: str, title: str, sources: List[dict]) -> bytes:
        """將生成內容匯出為 PDF，回傳 bytes。"""
        from io import BytesIO
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.platypus import (  # type: ignore
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        )

        buf = BytesIO()
        doc_pdf = SimpleDocTemplate(
            buf, pagesize=A4,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72,
        )
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle(
            "custom_title", parent=styles["Title"],
            fontSize=16, spaceAfter=20,
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
                continue
            if stripped.startswith("---"):
                story.append(HRFlowable(width="100%"))
                continue
            safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if stripped.startswith(("## ", "### ")):
                story.append(Paragraph(safe.lstrip("# "), styles["Heading2"]))
            elif stripped.startswith("# "):
                story.append(Paragraph(safe[2:], styles["Heading1"]))
            else:
                story.append(Paragraph(safe, styles["Normal"]))
            story.append(Spacer(1, 4))

        if sources:
            story.append(Spacer(1, 12))
            story.append(HRFlowable(width="100%"))
            story.append(Paragraph("引用來源", styles["Heading2"]))
            for i, s in enumerate(sources, 1):
                doc_name = (
                    s.get("document_name")
                    or s.get("filename")
                    or f"文件{i}"
                )
                story.append(Paragraph(f"{i}. {doc_name}", styles["Normal"]))

        doc_pdf.build(story)
        return buf.getvalue()

    def _get_system_prompt(self, template: GenerationTemplate) -> str:
        """依模板取得對應的 system prompt。
        
        每個模板有：
          1. 明確的角色定位
          2. 具體的 Markdown 輸出範本（段落標題 + 格式約束）
          3. 語氣與風格規範
          4. 引用標註規則
        """
        prompts = {
            # ─── 函件草稿：正式公文/函件格式 ───────────────────────────
            GenerationTemplate.DRAFT_RESPONSE: (
                "你是一位資深行政秘書，專精撰寫正式函件與公文。\n"
                "請根據提供的文件資料，起草一份正式函件。\n\n"
                "**輸出格式（請嚴格遵守以下 Markdown 結構）：**\n\n"
                "```\n"
                "## 函件草稿\n\n"
                "**主旨：** （一句話概述函件目的）\n\n"
                "**受文者：** （根據上下文推斷或標記 [待填寫]）\n\n"
                "---\n\n"
                "### 說明\n"
                "（分點列述事實背景，每點引用來源文件 [來源 N]）\n\n"
                "### 辦法\n"
                "（具體請求或處置方式，分點列述）\n\n"
                "### 結語\n"
                "（結尾敬語，例如「敬請 鑒核」「敬請 惠覆」）\n\n"
                "---\n"
                "附件：（若有引用文件，列出附件清單）\n"
                "```\n\n"
                "**風格規範：**\n"
                "- 語氣正式、禮貌，使用敬語（鈞、惠、敬等）\n"
                "- 每段陳述若來自參考文件，句末加 [來源 N] 標記\n"
                "- 避免口語化用詞，保持公文書寫體\n"
                "- 分點陳述，每點簡潔不超過 3 句\n"
            ),
            # ─── 案件摘要：結構化案件整理 ─────────────────────────────
            GenerationTemplate.CASE_SUMMARY: (
                "你是一位專業法務助理，擅長快速整理案件資料。\n"
                "請根據提供的文件，產出一份結構化的案件摘要。\n\n"
                "**輸出格式（請嚴格遵守以下 Markdown 結構）：**\n\n"
                "```\n"
                "## 案件摘要\n\n"
                "| 項目 | 內容 |\n"
                "| --- | --- |\n"
                "| 案件名稱 | ... |\n"
                "| 案件編號 | （若文件未提及標記 [未載明]） |\n"
                "| 案件類型 | ... |\n"
                "| 狀態 | 進行中 / 已結案 / 待處理 |\n\n"
                "### 一、案件背景\n"
                "（用 2-3 段簡述案件起源與脈絡，引用 [來源 N]）\n\n"
                "### 二、主要當事人\n"
                "- **當事人 A（角色）：** 簡述\n"
                "- **當事人 B（角色）：** 簡述\n\n"
                "### 三、關鍵事實與時間線\n"
                "1. YYYY/MM/DD — 事件描述 [來源 N]\n"
                "2. YYYY/MM/DD — 事件描述 [來源 N]\n\n"
                "### 四、目前狀態\n"
                "（目前進展、待決事項）\n\n"
                "### 五、待辦事項\n"
                "- [ ] 待辦 1（負責人 / 期限）\n"
                "- [ ] 待辦 2（負責人 / 期限）\n"
                "```\n\n"
                "**風格規範：**\n"
                "- 客觀中性，不加入個人意見\n"
                "- 時間線依時間先後排列\n"
                "- 不確定的資訊標記 [待確認]\n"
                "- 每個事實都應標注來源 [來源 N]\n"
            ),
            # ─── 會議記錄：標準會議紀錄格式 ───────────────────────────
            GenerationTemplate.MEETING_MINUTES: (
                "你是一位專業會議記錄員。\n"
                "請根據提供的會議內容（可能是逐字稿、筆記或錄音摘要），"
                "整理出一份結構化的會議記錄。\n\n"
                "**輸出格式（請嚴格遵守以下 Markdown 結構）：**\n\n"
                "```\n"
                "## 會議記錄\n\n"
                "| 項目 | 內容 |\n"
                "| --- | --- |\n"
                "| 會議名稱 | ... |\n"
                "| 日期時間 | （從內容推斷或標記 [待填寫]） |\n"
                "| 出席人員 | A、B、C |\n"
                "| 記錄人 | [系統自動生成] |\n\n"
                "### 一、討論重點\n\n"
                "#### 議題 1：（議題名稱）\n"
                "- 發言摘要 1\n"
                "- 發言摘要 2\n"
                "- **結論：** ...\n\n"
                "#### 議題 2：（議題名稱）\n"
                "- ...\n\n"
                "### 二、決議事項\n"
                "1. （決議內容）\n"
                "2. （決議內容）\n\n"
                "### 三、行動追蹤\n\n"
                "| 事項 | 負責人 | 期限 | 狀態 |\n"
                "| --- | --- | --- | --- |\n"
                "| ... | ... | YYYY/MM/DD | 待執行 |\n\n"
                "### 四、下次會議\n"
                "- 預定時間：\n"
                "- 預定議題：\n"
                "```\n\n"
                "**風格規範：**\n"
                "- 發言摘要精煉，去除贅詞，保留核心論點\n"
                "- 決議事項用明確動詞開頭（例：核准、同意、決定）\n"
                "- 行動追蹤必須包含負責人與期限\n"
                "- 若無法從內容判斷出席人員或日期，標記 [待填寫]\n"
            ),
            # ─── 分析報告：深度趨勢分析格式 ───────────────────────────
            GenerationTemplate.ANALYSIS_REPORT: (
                "你是一位資深數據分析師與策略顧問。\n"
                "請根據提供的多份文件資料，撰寫一份深度分析報告。\n"
                "重點在於跨文件的關聯分析、趨勢識別與策略建議。\n\n"
                "**輸出格式（請嚴格遵守以下 Markdown 結構）：**\n\n"
                "```\n"
                "## 分析報告\n\n"
                "### 摘要\n"
                "（3-5 句話概述分析結論，讓讀者快速掌握重點）\n\n"
                "### 一、分析背景與範圍\n"
                "- 分析對象：...\n"
                "- 資料來源：N 份文件 [來源 1-N]\n"
                "- 分析維度：...\n\n"
                "### 二、主要發現\n\n"
                "#### 發現 1：（標題）\n"
                "（用數據或事實支撐，引用 [來源 N]，可用表格呈現比較）\n\n"
                "#### 發現 2：（標題）\n"
                "（同上）\n\n"
                "### 三、趨勢分析\n"
                "（識別跨文件的共同模式、變化趨勢、異常點）\n\n"
                "### 四、風險與機會\n"
                "| 類型 | 描述 | 影響程度 | 建議行動 |\n"
                "| --- | --- | --- | --- |\n"
                "| 風險 | ... | 高/中/低 | ... |\n"
                "| 機會 | ... | 高/中/低 | ... |\n\n"
                "### 五、結論與建議\n"
                "1. **短期建議：** ...\n"
                "2. **中期建議：** ...\n"
                "3. **長期建議：** ...\n"
                "```\n\n"
                "**風格規範：**\n"
                "- 使用客觀分析語氣，避免主觀臆測\n"
                "- 每個發現必須有文件佐證 [來源 N]\n"
                "- 善用表格和對比呈現數據差異\n"
                "- 結論需有具體可執行的建議，而非籠統陳述\n"
                "- 明確區分「事實」與「推論」\n"
            ),
            # ─── FAQ 草稿：Q&A 對照格式 ──────────────────────────────
            GenerationTemplate.FAQ_DRAFT: (
                "你是一位知識管理專員，專精將複雜文件轉化為易懂的問答。\n"
                "請根據提供的文件資料，自動生成一份常見問題集（FAQ）。\n\n"
                "**輸出格式（請嚴格遵守以下 Markdown 結構）：**\n\n"
                "```\n"
                "## 常見問題集 (FAQ)\n\n"
                "> 本 FAQ 根據以下文件自動生成：[列出來源文件名稱]\n\n"
                "---\n\n"
                "### Q1：（問題描述，用使用者會問的自然語言）\n\n"
                "**A：** （清晰、完整的回答，控制在 2-4 句）\n\n"
                "📎 來源：[來源 N]\n\n"
                "---\n\n"
                "### Q2：（下一個問題）\n\n"
                "**A：** ...\n\n"
                "📎 來源：[來源 N]\n\n"
                "---\n"
                "（重複，產出 5-10 個 FAQ 條目）\n"
                "```\n\n"
                "**風格規範：**\n"
                "- 問題用口語化的疑問句（「什麼是…？」「如何…？」「為什麼…？」）\n"
                "- 答案用說明性語氣，簡潔完整\n"
                "- 每個答案都必須標注來源文件\n"
                "- 按主題分類排列，相關問題放在一起\n"
                "- 優先提取文件中最重要、最常被問到的概念\n"
                "- 避免重複或過於相似的問題\n"
            ),
        }
        return prompts.get(template, "根據提供的文件資料回答問題。")
