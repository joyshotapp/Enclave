"""
生成更新後的開發計畫 .docx
反映 Enclave 系統截至 Phase 13+ 的實際完成狀況
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 頁面邊距 ──────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── 輔助函式 ──────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(0x2e, 0x5e, 0x8e)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x4a, 0x7a, 0xa8)

def para(text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0')
        tcPr.append(shd)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ═══════════════════════════════════════════
#  封面
# ═══════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Enclave 企業 AI 知識大腦')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('On-Premise AI Knowledge Platform — 完整開發計畫（已完成版）')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x4a, 0x7a, 0xa8)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('版本：v2.0　｜　更新日期：2026-02-26　｜　狀態：Phase 13+ 全面完成\n').font.size = Pt(10)
meta.add_run('本文件反映系統實際完成狀況，供工程師獨立開發參考').font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════
#  一、產品定位
# ═══════════════════════════════════════════
h1('一、產品定位')

h2('1.1  核心價值主張')
para('Enclave 是一套完全地端部署的企業 AI 知識大腦，讓資料高度敏感的組織能在自己的伺服器上運行完整的 AI 問答、文件管理、內容生成等企業知識管理功能，資料永不外傳。')
bullet('目標客群：律師事務所、會計師事務所、醫療院所、政府機關、製造業、任何不適合雲端的專業服務業')
bullet('核心價值：資料完全不出境 × 主動索引文件 × AI 生成草稿 × 零 API 費用可選（全 Ollama）× 行動端隨時存取')
bullet('授權模式：地端一次性授權或年約，自行維護，不依賴第三方雲端')

h2('1.2  與原始 SaaS 版本的差異')
para('本系統由原 Phase 1–8 的多租戶 SaaS 架構演進而來，地端版本的主要調整方向：')
add_table(
    ['面向', 'SaaS 原始架構', '地端版實際實作'],
    [
        ['多租戶', '多企業共用，Row-Level Security', '保留 tenant_id 結構，預設單一組織模式；架構支援多租戶'],
        ['LLM 費用', '僅支援 OpenAI/Gemini 雲端 API', '支援 OpenAI / Gemini / Ollama 三種，可完全零 API 費用'],
        ['嵌入向量', '僅支援 Voyage AI 雲端嵌入', '支援 Voyage AI 或 Ollama bge-m3（零費用，已測試優先）'],
        ['SSO', '保留 SSO 模組（未移除）', '本機帳號 + JWT 為主，SSO 模組仍在但地端通常不啟用'],
        ['速率限制', '三層速率限制（多租戶版本）', '保留但簡化，支援單機資源保護'],
        ['文件解析', 'LlamaParse 優先', 'LlamaParse + 原生解析器 fallback + OCR（三層退化）'],
        ['部署方式', '雲端 SaaS 訂閱', 'Docker Compose 一鍵啟動，支援量產 + Linode 雲端部署'],
    ],
    [1.5, 2.0, 2.5]
)

# ═══════════════════════════════════════════
#  二、技術架構
# ═══════════════════════════════════════════
h1('二、技術架構')

h2('2.1  系統架構圖')
para('''
  [前端 Web]  React 19 + TypeScript + Vite（http://localhost:3001）
    [管理後台]  與前端同站點，採角色權限控管（superuser）
        │
        │ HTTPS / SSE
        ▼
  [Nginx]  反向代理 + SSL 終止
        │
        ▼
  [FastAPI]  Python 3.11  (http://localhost:8001 on-prem dev)
     ├── 認證：JWT HS256, OAuth2 Password Flow
     ├── 中介層：CORS, Rate Limit, Tenant 隔離, Admin IP Whitelist
     ├── 16 個 API 路由模組
     └── Celery 背景任務（文件向量化、Agent 批次）
         │
  ┌──────┴──────┐
  [Redis 7]    [PostgreSQL 15 + pgvector]
  任務佇列      多資料表 + tenant_id 隔離
  快取(:6380)  1024-dim 向量欄位 HNSW 索引 (:5435)
''', indent=1)

h2('2.2  核心技術選型')
add_table(
    ['元件', '技術', '說明'],
    [
        ['後端框架', 'FastAPI (Python 3.11)', 'async/await, Pydantic v2, SQLAlchemy 2'],
        ['前端框架', 'React 19 + TypeScript + Vite', 'React Router v7, TanStack Query, shadcn/ui'],
        ['資料庫', 'PostgreSQL 15 + pgvector', '向量相似度搜尋 (HNSW 索引), 1024-dim'],
        ['任務佇列', 'Celery + Redis 7', '文件向量化非同步執行'],
        ['LLM（預設）', 'Gemini (via OpenAI-compatible API)', 'LLM_PROVIDER=gemini, 可改 openai 或 ollama'],
        ['LLM（零費用）', 'Ollama (本機)', '支援 llama3-taiwan, ministral, gpt-oss 等'],
        ['嵌入向量（預設）', 'Ollama bge-m3', '1024-dim, EMBEDDING_PROVIDER=ollama (零費用)'],
        ['嵌入向量（備選）', 'Voyage AI voyage-4-lite', '1024-dim, 雲端付費方案'],
        ['文件解析', 'LlamaParse + PyMuPDF + pytesseract', '三層退化：LlamaParse → 原生 → OCR'],
        ['中文分詞', 'jieba', 'BM25 關鍵詞搜尋用'],
        ['行動端', 'React Native (Expo)', 'iOS + Android 雙平台'],
        ['容器化', 'Docker Compose', '5 個服務：web, worker, db, redis, frontend'],
    ],
    [1.8, 2.0, 2.5]
)

h2('2.3  本機開發埠號（on-prem 實際運行）')
add_table(
    ['服務', '容器名稱', '對外埠', '說明'],
    [
        ['FastAPI 後端', 'enclave-web-1', '8001', '原始碼掛載，改動即時生效'],
        ['React 前端', 'enclave-frontend-1', '3001', 'Nginx 靜態服務'],
        ['PostgreSQL', 'enclave-db-1', '5435', '用 5435 避免與本機衝突'],
        ['Redis', 'enclave-redis-1', '6380', '任務佇列 + 快取'],
        ['Celery Worker', 'enclave-worker-1', '—', '無對外埠，背景向量化任務'],
    ],
    [1.5, 1.8, 1.0, 2.0]
)

# ═══════════════════════════════════════════
#  三、功能模組總覽
# ═══════════════════════════════════════════
h1('三、功能模組總覽（全部已完成）')

add_table(
    ['模組', 'API 前綴', '說明'],
    [
        ['AI 問答', '/chat', 'RAG 問答，SSE 串流，多輪對話（代名詞追蹤），回饋評分'],
        ['內容生成', '/generate', '5 種模板 RAG 增強生成，SSE 串流，Word/PDF 匯出'],
        ['報告管理', '/generate/reports', '生成報告持久化，釘選，分頁搜尋，CRUD'],
        ['文件管理', '/documents', '23 格式上傳，Celery 自動解析向量化，狀態追蹤'],
        ['Agent 自動索引', '/agent', '資料夾監控，AI 分類提案，人工審核佇列，批次排程'],
        ['KB 維護', '/kb-maintenance', '版本管理，健康度儀表板，缺口偵測，備份還原，分類管理'],
        ['知識庫搜尋', '/kb', '混合語意搜尋，top_k，重排序開關'],
        ['問答分析', '/chat/analytics', '摘要，趨勢，熱門問題，知識缺口（admin+）'],
        ['公司管理', '/company', '成員邀請，用量統計，配額'],
        ['管理後台', '/admin', '儀表板，使用者管理，系統健康（superuser）'],
        ['部門管理', '/departments', '部門 CRUD，樹狀結構，內聯編輯'],
        ['稽核日誌', '/audit', '操作日誌，用量記錄，成本統計，CSV 匯出'],
        ['行動端 API', '/mobile', 'JWT 刷新，推播 token，安全事件'],
        ['功能旗標', '/feature-flags', '功能開關管理（superuser，per-tenant）'],
    ],
    [1.5, 1.8, 3.1]
)

# ═══════════════════════════════════════════
#  四、RAG 檢索引擎設計
# ═══════════════════════════════════════════
h1('四、RAG 檢索引擎設計')

h2('4.1  混合檢索流程')
para('問答時完整流程：')
bullet('1. 多輪代名詞偵測：偵測「它」「這個」「其中」「這些」等指示詞，改寫為完整問題（chat_orchestrator.py）')
bullet('2. 語意搜尋：Ollama bge-m3 嵌入 + pgvector cosine 相似度')
bullet('3. BM25 關鍵詞搜尋：jieba 中文分詞')
bullet('4. RRF 融合排序：Reciprocal Rank Fusion 合併兩路結果')
bullet('5. 本地重排序（_local_rerank）：jieba 關鍵字重疊(35%) + 實體比對(25%) + RRF 分數(40%)')
bullet('6. LLM 生成回答：Gemini / OpenAI / Ollama（可設定）')
bullet('7. SSE 串流回傳 + 來源引用')

h2('4.2  關鍵參數')
add_table(
    ['參數', '預設值', 'config 變數'],
    [
        ['分塊大小', '1000 tokens', 'CHUNK_SIZE'],
        ['分塊重疊', '150 tokens', 'CHUNK_OVERLAP'],
        ['向量維度', '1024-dim', 'EMBEDDING_DIMENSION'],
        ['檢索數量', '5', 'RETRIEVAL_TOP_K'],
        ['最低分數', '0.0（不過濾）', 'RETRIEVAL_MIN_SCORE'],
        ['快取 TTL', '300 秒', 'RETRIEVAL_CACHE_TTL'],
        ['HNSW m', '16', 'db migration 設定'],
        ['HNSW ef_construction', '64', 'db migration 設定'],
    ],
    [1.5, 2.0, 2.5]
)

h2('4.3  Chunk 設計重點')
para('chunk 儲存時第一個 chunk 及短 chunk（<800 字）會加入 【文件名稱】 前綴，讓向量附帶文件身份資訊，改善同類型文件的辨識準確度（在 document_tasks.py 實作）。')

h2('4.4  LLM 提供者切換')
add_table(
    ['LLM_PROVIDER', '說明', '必要環境變數'],
    [
        ['gemini（預設）', 'Google Gemini API，使用 OpenAI 相容介面', 'GEMINI_API_KEY, GEMINI_MODEL'],
        ['openai', 'OpenAI API', 'OPENAI_API_KEY, OPENAI_MODEL'],
        ['ollama', '本機 Ollama，零費用', 'OLLAMA_BASE_URL（Docker 內用 host.docker.internal:11434）, OLLAMA_MODEL'],
    ],
    [1.5, 2.5, 2.5]
)
para('注意：Docker 容器內 OLLAMA_BASE_URL 必須用 http://host.docker.internal:11434，不能用 localhost（localhost 在容器內指向容器自身）。嵌入向量有獨立的 OLLAMA_EMBED_URL 設定，預設已正確設為 host.docker.internal。', bold=False)

# ═══════════════════════════════════════════
#  五、角色與權限
# ═══════════════════════════════════════════
h1('五、角色與權限系統')

h2('5.1  角色層級')
add_table(
    ['角色', '說明'],
    [
        ['superuser', '系統管理員，跨所有租戶，僅用於系統初始化與維運'],
        ['owner', '組織擁有者，可管理本租戶所有資源'],
        ['admin', '管理員，可管理使用者與文件，看所有稽核日誌'],
        ['hr', '人資 / 業務，可問答、生成、審核文件'],
        ['employee', '一般員工，可問答與查閱文件'],
        ['viewer', '唯讀，可查看與使用一般端點'],
    ],
    [1.5, 4.5]
)

h2('5.2  關鍵端點最低角色')
add_table(
    ['操作', '最低角色'],
    [
        ['AI 問答 / 串流 / 歷史', 'viewer'],
        ['文件列表 / 詳情', 'viewer（依部門範圍）'],
        ['文件上傳 / 刪除', 'hr'],
        ['內容生成', 'viewer'],
        ['Agent 設定 / 審核佇列', 'admin'],
        ['問答分析 / 用量統計（租戶層級）', 'admin'],
        ['使用者建立', 'admin'],
        ['組織更新', 'superuser'],
        ['功能旗標管理', 'superuser'],
    ],
    [3.0, 3.0]
)

# ═══════════════════════════════════════════
#  六、資料模型
# ═══════════════════════════════════════════
h1('六、主要資料表')

add_table(
    ['資料表', '模型類別', '說明'],
    [
        ['tenants', 'Tenant', '組織 / 租戶'],
        ['users', 'User', '使用者（角色、部門、superuser）'],
        ['departments', 'Department', '部門（樹狀）'],
        ['documents', 'Document', '文件元資料，status: pending→processing→completed/failed'],
        ['documentchunks', 'DocumentChunk', '文件分塊 + 1024-dim 向量（pgvector HNSW）'],
        ['conversations', 'Conversation', '對話'],
        ['messages', 'Message', '對話訊息（user/assistant）'],
        ['retrievaltraces', 'RetrievalTrace', 'RAG 檢索追蹤，sources_json'],
        ['auditlogs', 'AuditLog', '操作日誌'],
        ['usagerecords', 'UsageRecord', '用量記錄（tokens/latency/cost）'],
        ['generated_reports', 'GeneratedReport', '內容生成報告持久化'],
        ['watch_folders', 'WatchFolder', 'Agent 監控資料夾'],
        ['review_items', 'ReviewItem', '待審核文件佇列，status: pending→approved/rejected'],
        ['documentversions', 'DocumentVersion', '文件版本歷史'],
        ['categories', 'Category', '文件分類（樹狀）'],
        ['kbbackups', 'KBBackup', 'KB 備份任務'],
        ['knowledgegaps', 'KnowledgeGap', '知識缺口記錄'],
        ['integrityreports', 'IntegrityReport', '完整性掃描報告'],
        ['feature_flags', 'FeatureFlag', '功能旗標'],
        ['chat_feedbacks', 'ChatFeedback', '回饋評分'],
    ],
    [1.8, 1.8, 2.8]
)

# ═══════════════════════════════════════════
#  七、開發階段狀態
# ═══════════════════════════════════════════
h1('七、開發階段完成狀態')

h2('7.1  階段總覽')
add_table(
    ['Phase', '名稱', '核心交付', '狀態'],
    [
        ['1–2', '帳號管理 + 認證', '多租戶 + JWT + 文件管線 + AI 問答', '✅ 完成'],
        ['3–4', '管理功能 + 生產化', '配額 + 分析 + CI/CD + 監控', '✅ 完成'],
        ['5–6', 'UX + AI 引擎升級', 'SSE 串流 + LlamaParse + jieba + HyDE', '✅ 完成'],
        ['7–8', '對話體驗 + 安全加固', 'SSE 串流 + Markdown + 來源面板 + Docker', '✅ 完成'],
        ['9', '地端轉型基礎', 'LLM 抽象層 + Ollama + 地端安裝包 + 管理後台精簡', '✅ 完成'],
        ['10', '主動索引 Agent', 'File Watcher + AI 分類 + 人工審核 UI + 排程批次', '✅ 完成'],
        ['11', '內容生成引擎', '5 模板 RAG 生成 + SSE + Word/PDF 匯出', '✅ 完成'],
        ['11-2', '報告持久化', '報告 CRUD + 日期分組 + 釘選', '✅ 完成'],
        ['12', '行動端 App', 'React Native iOS+Android + 問答 + 上傳 + 生成', '✅ 完成'],
        ['13', 'KB 主動維護', '版本管理 + 健康度儀表板 + 缺口偵測 + 備份 + 分類', '✅ 完成'],
        ['13+', 'UX 優化', '側欄重整 + 頁面合併 + 分析整合 + RAG 品質優化', '✅ 完成'],
        ['14', '進階 Agent', 'Rule Engine + Digest + Multi-step', '🔲 計畫中'],
        ['15', '協作功能', '標註 + 分享 + 知識圖譜', '🔲 計畫中'],
    ],
    [0.8, 1.6, 3.0, 1.0]
)

h2('7.2  Phase 9 地端轉型——實際完成情況')
para('原計畫：完全移除多租戶架構、SSO、三層速率限制。\n實際執行：採保守策略，架構層面保留但在應用層預設單一組織模式。原因：')
bullet('保留 tenant_id 讓架構具備未來擴展性（集團多部門、多子公司）')
bullet('SSO 模組保留在程式庫但地端通常不啟用')
bullet('速率限制保留但參數調整，用於保護單機資源')
bullet('新增 INTERNAL_LLM_PROVIDER 設定，分類/改寫等輕任務走 Ollama，問答走設定的主 LLM')

h2('7.3  Phase 13+ RAG 品質優化——實際成果')
para('在原計畫 Phase 13 框架外增加的品質優化工作（2026-02-25 完成）：')
add_table(
    ['優化項目', '實作位置', '效果'],
    [
        ['本地重排序器（_local_rerank）', 'kb_retrieval.py', '取代 Voyage Rerank，零 API 費用，兼顧效果'],
        ['多輪代名詞擴充', 'chat_orchestrator.py _CONTEXT_PRONOUNS', '多輪追問正確率 0% → 100%'],
        ['Chunk 文件名稱前綴', 'document_tasks.py', '改善同類型文件的向量辨識'],
        ['結構化問答健診修正', 'structured_answers.py', '移除錯誤的硬查詢路由'],
        ['整體 RAG 品質', '（17 測試案例）', '71.1% → 75.5%，整體評級「優秀」'],
    ],
    [2.0, 2.0, 2.5]
)

# ═══════════════════════════════════════════
#  八、快速啟動
# ═══════════════════════════════════════════
h1('八、快速啟動')

h2('8.1  前置需求')
bullet('Docker Engine 24+ & Docker Compose v2')
bullet('至少 8GB RAM（全 Ollama 模式需要 12GB+）')
bullet('Python 3.11+（本機開發；Docker 環境已內建）')
bullet('Node.js 20+（前端開發）')
bullet('Ollama（如需零費用模式）：安裝後執行 ollama pull bge-m3')

h2('8.2  啟動步驟')
para('步驟 1：複製環境變數')
para('cp .env.example .env', indent=1)
para('步驟 2：填寫必要變數（最少設定）', bold=True)
para('''SECRET_KEY=<執行 python -c "import secrets; print(secrets.token_urlsafe(48))" 產生>
GEMINI_API_KEY=<your-gemini-key>    ← 或改用 LLM_PROVIDER=ollama 零費用
ORGANIZATION_NAME=我的公司''', indent=1)
para('步驟 3：啟動所有服務')
para('docker compose up -d', indent=1)
para('步驟 4：初始化資料庫')
para('''docker compose exec web alembic upgrade head
docker compose exec web python scripts/initial_data.py''', indent=1)
para('步驟 5：開啟系統')
add_table(
    ['入口', '網址'],
    [
        ['前端 Web', 'http://localhost:3001'],
        ['API 文件（Swagger）', 'http://localhost:8001/docs'],
        ['API 健康檢查', 'http://localhost:8001/health'],
    ],
    [2.0, 4.0]
)
para('預設登入帳號：admin@example.com / admin123（請在生產環境立即修改）')

h2('8.3  零費用 Ollama 模式')
para('如要完全不使用任何付費 API，在 .env 設定：')
para('''EMBEDDING_PROVIDER=ollama
OLLAMA_EMBED_MODEL=bge-m3
OLLAMA_EMBED_URL=http://host.docker.internal:11434

LLM_PROVIDER=ollama
OLLAMA_MODEL=kenneth85/llama-3-taiwan:8b-instruct-q8_0
OLLAMA_BASE_URL=http://host.docker.internal:11434''', indent=1)
para('注意：Docker 容器內必須用 host.docker.internal，Ollama 要在宿主機上運行（非容器內）。')

h2('8.4  Alembic migration 鏈')
add_table(
    ['Migration ID', '說明'],
    [
        ['450ae450e023', '初始 Schema（所有核心資料表）'],
        ['84a829cdf24b', '部門 + 功能權限'],
        ['f7859742ce5d', '租戶 SSO 設定'],
        ['eb7fe95812e8', '功能旗標'],
        ['a1b2c3d4e5f6', 'pgvector 嵌入欄位'],
        ['c7c9c43b1a3d', 'Custom Domains'],
        ['d1e2f3a4b5c6', 'Phase 7/10/13 資料表 ← HEAD'],
    ],
    [2.5, 4.0]
)

# ═══════════════════════════════════════════
#  九、環境變數
# ═══════════════════════════════════════════
h1('九、關鍵環境變數')

h2('9.1  核心必填')
add_table(
    ['變數', '說明', '範例值'],
    [
        ['SECRET_KEY', 'JWT 簽名密鑰，≥32 字元', 'python -c "import secrets; print(secrets.token_urlsafe(48))"'],
        ['ORGANIZATION_NAME', '顯示在 UI 的組織名稱', '泰宇科技'],
        ['APP_ENV', '環境類型', 'development / production'],
    ],
    [2.0, 2.0, 2.5]
)

h2('9.2  LLM / 嵌入')
add_table(
    ['變數', '預設值', '說明'],
    [
        ['LLM_PROVIDER', 'gemini', 'openai / gemini / ollama'],
        ['GEMINI_API_KEY', '—', 'Gemini 模式必填'],
        ['GEMINI_MODEL', 'gemini-3-flash-preview', '⚠️ 勿隨意更改此值'],
        ['OPENAI_API_KEY', '—', 'OpenAI 模式必填'],
        ['OPENAI_MODEL', 'gpt-4o-mini', '—'],
        ['OLLAMA_BASE_URL', 'http://localhost:11434', 'Docker 內改 host.docker.internal:11434'],
        ['OLLAMA_MODEL', 'llama3.2', '建議 kenneth85/llama-3-taiwan:8b-instruct-q8_0'],
        ['EMBEDDING_PROVIDER', 'ollama', 'ollama（零費用）或 voyage'],
        ['OLLAMA_EMBED_URL', 'http://host.docker.internal:11434', 'Docker 內嵌入向量專用 URL'],
        ['OLLAMA_EMBED_MODEL', 'bge-m3', '1024-dim 中英文嵌入'],
        ['VOYAGE_API_KEY', '—', 'EMBEDDING_PROVIDER=voyage 時才需要'],
        ['LLAMAPARSE_API_KEY', '—', '選填，停用則用原生解析器'],
        ['INTERNAL_LLM_PROVIDER', 'ollama', '分類/改寫輕任務用（不消耗主 LLM 額度）'],
    ],
    [2.2, 2.2, 2.0]
)

h2('9.3  檢索參數')
add_table(
    ['變數', '預設值', '說明'],
    [
        ['RETRIEVAL_MODE', 'hybrid', 'semantic / keyword / hybrid'],
        ['RETRIEVAL_TOP_K', '5', '預設檢索數量'],
        ['RETRIEVAL_MIN_SCORE', '0.0', '最低相似度閾值（0.0=不過濾）'],
        ['RETRIEVAL_RERANK', 'true', '啟用重排序（無 Voyage key 時自動用本地重排序）'],
        ['CHUNK_SIZE', '1000', '分塊大小（tokens）'],
        ['CHUNK_OVERLAP', '150', '分塊重疊（tokens）'],
    ],
    [2.2, 1.5, 2.8]
)

# ═══════════════════════════════════════════
#  十、專案目錄結構
# ═══════════════════════════════════════════
h1('十、專案目錄結構')

para('''Enclave/
├── app/
│   ├── api/
│   │   ├── deps.py                  # 依賴注入（get_db, get_current_user）
│   │   ├── deps_permissions.py      # RBAC 權限（require_*）
│   │   └── v1/endpoints/            # 16 個端點模組
│   │       ├── auth.py  users.py  tenants.py  admin.py  company.py
│   │       ├── documents.py  kb.py  chat.py  generate.py  reports.py
│   │       ├── agent.py  kb_maintenance.py  audit.py  analytics.py
│   │       ├── departments.py  feature_flags.py  mobile.py
│   ├── services/
│   │   ├── document_parser.py       # 23 格式文件解析（三層退化）
│   │   ├── kb_retrieval.py          # 混合檢索 + 本地重排序
│   │   ├── chat_orchestrator.py     # SSE 串流問答 + 多輪上下文
│   │   ├── content_generator.py     # RAG 增強生成 + 匯出
│   │   ├── llm_client.py            # OpenAI/Gemini/Ollama 統一介面
│   │   └── structured_answers.py   # 結構化回答（薪資/健康等）
│   ├── agent/
│   │   ├── file_watcher.py          # watchdog 資料夾監控
│   │   ├── classifier.py            # AI 文件分類
│   │   ├── review_queue.py          # 審核佇列狀態機
│   │   └── scheduler.py            # 排程批次
│   ├── tasks/
│   │   └── document_tasks.py        # Celery 向量化任務（含 chunk prefix）
│   ├── models/                      # SQLAlchemy ORM 模型
│   ├── schemas/                     # Pydantic v2 Schema
│   ├── crud/                        # 資料存取層
│   ├── middleware/                  # CORS, Rate Limit, Tenant
│   └── config.py                   # 所有設定（含預設值）
├── frontend/                        # React 19 + Vite 主前端
│   └── src/
│       ├── pages/                   # 15+ 頁面
│       ├── components/Layout.tsx    # 側欄導覽（4 分組，角色過濾）
│       ├── api.ts                   # API 客戶端
│       └── types.ts
├── admin_service/                   # 獨立管理後台服務
│   └── main.py
├── admin-frontend/                  # 管理後台前端
├── alembic/versions/                # Alembic migration
├── docker-compose.yml               # 開發環境（掛載原始碼）
├── docker-compose.prod.yml          # 生產環境
├── nginx/                           # Nginx 設定
├── monitoring/                      # Prometheus + Grafana
├── scripts/
│   ├── initial_data.py             # 初始化 superuser + Demo Tenant
│   ├── batch_upload.py             # 批量上傳測試文件
│   ├── check_retrieval_quality.py  # RAG 品質驗證（17 測試案例）
│   ├── generate_test_data.py       # 生成 100+ 測試文件
│   └── ...
└── test-data/                       # 101 份測試文件（各類型）''', indent=1)

# ═══════════════════════════════════════════
#  十一、工作流程
# ═══════════════════════════════════════════
h1('十一、關鍵工作流程')

h2('11.1  文件上傳到可問答')
para('''使用者上傳文件
  → POST /api/v1/documents/upload（回傳 document.id + status="pending"）
  → Celery 任務 process_document_task
      1. 文件解析（LlamaParse / 原生 / OCR 三層退化）
      2. Markdown 分塊（chunk_size=1000, overlap=150）
      3. Chunk 加入【文件名稱】前綴（第一個 chunk 及短 chunk）
      4. Ollama bge-m3 嵌入（1024-dim）
      5. pgvector 寫入（HNSW 索引）
      6. document.status = "completed"
  → 可問答 ✅''', indent=1)

h2('11.2  Agent 自動索引')
para('''磁碟資料夾（AGENT_WATCH_FOLDERS）
  → file_watcher（watchdog）偵測新檔案
  → review_queue.enqueue（ReviewItem status=pending）
  → 管理員在 /agent/review 審核
      [核准] → 觸發向量化任務 → completed
      [拒絕] → 標記拒絕
      [批量核准] → 高信心文件一鍵處理
  排程批次：每日深夜 AGENT_BATCH_HOUR（預設凌晨 2 點）''', indent=1)

h2('11.3  SSE 串流問答請求格式')
para('''POST /api/v1/chat/chat/stream
  Authorization: Bearer <token>
  Content-Type: application/json

  {
    "question": "陳建宏的本薪是多少？",
    "conversation_id": "uuid（選填，延續多輪對話）",
    "top_k": 5
  }

  SSE 事件格式：
  data: {"type": "status",  "content": "正在搜尋文件..."}
  data: {"type": "token",   "content": "根據您的薪資條..."}
  data: {"type": "done",    "message_id": "uuid", "conversation_id": "uuid",
                            "sources": [{"title": "202602-E001-陳建宏-薪資條.txt", ...}]}''', indent=1)

# ═══════════════════════════════════════════
#  十二、測試與品質驗證
# ═══════════════════════════════════════════
h1('十二、測試與品質驗證')

h2('12.1  RAG 品質基準（截至 2026-02-25）')
add_table(
    ['類別', '測試案例數', '平均命中率'],
    [
        ['薪資制度', '2', '62%'],
        ['勞動法規', '4', '83%'],
        ['請假規定', '2', '71%'],
        ['採購報帳', '2', '84%'],
        ['績效考核', '2', '75%'],
        ['組織人力', '2', '62%'],
        ['健康管理', '1', '100%'],
        ['教育訓練', '1', '75%'],
        ['多輪對話', '1+1 追問', '50% + 100%'],
        ['整體', '17', '75.5%（優秀）'],
    ],
    [2.0, 1.8, 2.0]
)

h2('12.2  品質測試工具')
bullet('scripts/check_retrieval_quality.py — 17 個測試案例，自動計算關鍵字命中率')
bullet('scripts/generate_test_data.py — 生成 101 份台灣繁體中文企業測試文件')
bullet('scripts/batch_upload.py — 批量上傳並等待 Celery 處理完成')
bullet('執行：python -X utf8 scripts/check_retrieval_quality.py')

h2('12.3  常見問題排查')
add_table(
    ['症狀', '原因', '解法'],
    [
        ['嵌入失敗 / 文件卡在 processing', 'Ollama 未啟動或 URL 錯誤', '確認 ollama serve 運行中；Docker 內用 host.docker.internal'],
        ['問答回答錯誤員工資訊', 'Redis 快取殘留', 'docker compose exec redis redis-cli FLUSHALL'],
        ['多輪追問無法找到上文', '代名詞未觸發上下文改寫', '確認問句含有「其中」「這些」等指示詞，或直接把問題說完整'],
        ['Docker 容器無法呼叫 Ollama LLM', 'OLLAMA_BASE_URL=localhost（錯誤）', '改為 http://host.docker.internal:11434'],
        ['文件外鍵錯誤無法刪除', 'documentversions 參照', '先 DELETE FROM documentversions 再刪 documents'],
    ],
    [1.8, 1.8, 2.2]
)

# ═══════════════════════════════════════════
#  十三、給獨立開發工程師的建議
# ═══════════════════════════════════════════
h1('十三、給獨立開發工程師的建議')

h2('13.1  建議開發順序')
para('如從零開始獨立開發本系統，建議順序同 Phase 編號，關鍵依賴：')
bullet('Phase 1–8（基礎）必須先完成，所有後續 Phase 都依賴這個骨架')
bullet('Phase 9（LLM 抽象層）是地端版的關鍵轉換點，完成後才能無縫切換 LLM')
bullet('Phase 10（Agent）與 Phase 11（生成引擎）在後端無強依賴，可並行開發')
bullet('Phase 12（行動 App）依賴 Phase 9 API 架構穩定後才開發')
bullet('Phase 13（KB 維護）建議資料量充足後再做，健康度指標才有意義')

h2('13.2  容易踩到的坑')
bullet('Ollama URL：Docker 容器內永遠用 host.docker.internal，不是 localhost')
bullet('pgvector 維度：建立 embedding 欄位時就決定維度（1024），之後改維度需要 migration 重建欄位')
bullet('多輪對話：代名詞列表要夠完整（「其中」「這些」「那些」「上面」等），否則追問查不到上文')
bullet('Chunk 去重：同一份文件重複上傳不應重複向量化，用 chunk_hash（SHA256）做去重')
bullet('SSE 串流：前端要用 EventSource 或 fetch+ReadableStream，不能用普通 axios')
bullet('Celery worker：文件向量化是非同步的，上傳後要輪詢 status 欄位，不是立即可問答')
bullet('DB 外鍵順序：刪 documents 前要先刪 documentversions、documentchunks')

h2('13.3  效能基準（本機 on-prem 環境）')
add_table(
    ['操作', '平均耗時', '備註'],
    [
        ['文件向量化（1MB MD）', '~3-8 秒', 'Celery worker + Ollama bge-m3'],
        ['問答回應（SSE 串流）', '~6-12 秒', 'Gemini + hybrid 檢索 + 本地重排序'],
        ['批量上傳 101 份文件', '~20 秒全部完成', 'Celery 並行處理'],
        ['Redis 快取命中', '<0.1 秒', '相同問題第二次查詢'],
    ],
    [2.0, 1.5, 3.0]
)

# ═══════════════════════════════════════════
#  尾頁
# ═══════════════════════════════════════════
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.add_run('Enclave 開發計畫 v2.0　｜　2026-02-26　｜　私有專案，未經授權不得複製或分發').font.size = Pt(9)

# ── 輸出 ──────────────────────────────────────
output = r'C:\Users\User\Desktop\Enclave\開發計畫_企業AI知識大腦_地端版.docx'
doc.save(output)
print(f'✅ 已儲存：{output}')
